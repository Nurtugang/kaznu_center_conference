import os
from io import BytesIO
from docx import Document
from django.core.management.base import BaseCommand, CommandError
from django.core.files.base import ContentFile
from django.contrib.auth import get_user_model
from conferences.models import Conference, Submission, SubmissionVersion

User = get_user_model()

class Command(BaseCommand):
    help = 'Создает 5 тестовых заявок для единственной конференции в системе'

    def handle(self, *args, **kwargs):
        # 1. Проверка количества конференций
        conf_count = Conference.objects.count()
        
        if conf_count == 0:
            raise CommandError("Ошибка: В системе нет ни одной конференции. Сначала создайте конференцию через админку.")
        
        if conf_count > 1:
            raise CommandError(f"Ошибка: В системе найдено {conf_count} конф. Скрипт рассчитан на работу только с ОДНОЙ конференцией.")

        conf = Conference.objects.first()
        self.stdout.write(self.style.SUCCESS(f"Работаем с конференцией: {conf.title} (ID: {conf.id})"))

        for i in range(1, 6):
            # 2. Создаем или получаем пользователя
            username = f'author_{i}'
            user, created = User.objects.get_or_create(
                username=username,
                defaults={
                    'email': f'{username}@example.com', 
                    'role': 'author',
                    'first_name': f'Имя_{i}',
                    'last_name': f'Фамилия_{i}',
                    'organization': 'КазНУ им. аль-Фараби'
                }
            )
            if created:
                user.set_password('password123')
                user.save()

            # 3. Генерируем Word-файл (.docx) в оперативной памяти
            doc = Document()
            doc.add_heading(f'Научное исследование №{i}', 0)
            doc.add_paragraph(f'Автор: {user.get_full_name()}')
            doc.add_paragraph('Аннотация: Это тестовое содержание статьи для проверки системы загрузки и обработки файлов.')
            
            # Сохраняем ворд в буфер
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_filename = f'paper_id{i}_draft.docx'
            docx_content = ContentFile(docx_buffer.getvalue(), name=docx_filename)

            # 4. Создаем заявку (Submission)
            submission, sub_created = Submission.objects.get_or_create(
                user=user,
                conference=conf,
                defaults={
                    'title': f'Тема научной работы №{i}',
                    'authors_list': f'{user.last_name} {user.first_name}',
                    'status': 'under_review'
                }
            )

            # 5. Создаем версию (SubmissionVersion) и сохраняем туда Word
            version = SubmissionVersion.objects.create(
                submission=submission,
                version_number=1,
                author_comment="Начальная версия в формате DOCX"
            )
            version.file.save(docx_filename, docx_content, save=True)

            self.stdout.write(f' - Заявка от {username} создана.')

        self.stdout.write(self.style.SUCCESS('Готово! 5 заявок успешно добавлены.'))